import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx2pdf import convert


def extract_text_from_beginning_to_keyword(text, end_keyword):
    pattern = re.compile('(.*?)' + re.escape(end_keyword), re.S)
    matches = pattern.search(text)
    return matches.group(1)


def set_spacing_for_heading(heading):
    for run in heading.runs:
        run.font.size = Pt(14)
    paragraph_format = heading.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def set_no_spacing_paragraph(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0


def create_cover_letter(cover_letter_chatgpt, user):
    doc = Document()
    subheading = f'{user[1]} {user[2]}\n{user[3]}, {user[5]}\n{user[6]}\n{user[7]}'
    subheading_paragraph = doc.add_paragraph()
    subheading_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = subheading_paragraph.add_run(subheading, 0)
    run.bold = False
    run.font.size = Pt(11)

    text_heading = doc.add_heading('', level=1)
    set_spacing_for_heading(text_heading)
    text_paragraph = doc.add_paragraph(extract_text_from_beginning_to_keyword(cover_letter_chatgpt, "END"))
    set_no_spacing_paragraph(text_paragraph)

    # Save the document
    doc.save('cover_letter_created.docx')
    convert('cover_letter_created.docx', 'cover_letter_created.pdf')
