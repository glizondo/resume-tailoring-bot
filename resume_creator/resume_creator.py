import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert


def extract_text_between_keywords(text, start_keyword, end_keyword):
    pattern = re.compile(re.escape(start_keyword) + '(.*?)' + re.escape(end_keyword), re.S)
    matches = pattern.search(text)
    return matches.group(1)


def set_spacing_for_heading(heading):
    for run in heading.runs:
        run.font.size = Pt(14)
    paragraph_format = heading.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

def set_no_spacing_paragraph(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0

def create_resume(resume_chat_gpt, user):
    doc = Document()
    heading = doc.add_heading(f'{user[1]} {user[2]}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subheading = f'{user[3]}, {user[5]} - {user[6]} - {user[7]} - {user[8]}'
    subheading_paragraph = doc.add_paragraph()
    subheading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subheading_paragraph.add_run(subheading, 0)
    run.bold = True
    run.font.size = Pt(14)

    education = str(extract_text_between_keywords(resume_chat_gpt, "Education:", "Work Experiences:"))
    work_experience = str(extract_text_between_keywords(resume_chat_gpt, "Work Experiences:", "Skills:"))
    skills = str(extract_text_between_keywords(resume_chat_gpt, "Skills:", "END"))

    # Education
    education_heading = doc.add_heading('Education:', level=1)
    set_spacing_for_heading(education_heading)
    education_paragraph = doc.add_paragraph(education)
    set_no_spacing_paragraph(education_paragraph)

    # Experience
    experience_heading = doc.add_heading('Work Experience:', level=1)
    set_spacing_for_heading(experience_heading)
    for line in work_experience.split('\n'):
        if line.startswith('-'):
            paragraph = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            set_no_spacing_paragraph(paragraph)
        else:
            paragraph = doc.add_paragraph(line)
            set_no_spacing_paragraph(paragraph)



    # Skills
    skills_heading = doc.add_heading('Skills:', level=1)
    set_spacing_for_heading(skills_heading)
    skills_paragraph = doc.add_paragraph(skills)
    set_no_spacing_paragraph(skills_paragraph)

    # Save the document
    doc.save('resume_created.docx')
    convert('resume_created.docx', 'resume_created.pdf')
